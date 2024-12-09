from docx import Document
from docx.shared import Inches
import os
import pandas as pd
from docx.enum.section import WD_ORIENTATION
import datetime 

def create_document(dictionary: dict):

    document = Document()
    document.add_heading(f"Izvještaj prosjeka čišćenja {datetime.datetime.now()}")
    
    for key, value in dictionary.items():
        if isinstance(value, pd.Series):
            value = value.to_frame(name="Values").reset_index()
        document.add_heading(key, level=2)

        table = document.add_table(rows=1, cols=len(value.columns))
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(value.columns):
            hdr_cells[i].text = column_name

        for _, row in value.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)

    document.save("test.docx")